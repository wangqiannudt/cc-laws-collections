#!/usr/bin/env python3
"""导入内部法规到数据库"""
import sys
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from datetime import date
from docx import Document

from app.database import SessionLocal
from app.models.law import Law, create_law


def import_docx(file_path: str, category: str = "内部法规"):
    """
    导入 docx 格式的内部法规

    Args:
        file_path: docx 文件路径
        category: 分类名称
    """
    db = SessionLocal()

    try:
        # 读取 docx 文件
        doc = Document(file_path)

        # 提取标题（通常是第一个段落）
        title = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                title = text
                break

        if not title:
            print("错误: 无法从文档中提取标题")
            return

        # 提取正文内容
        content_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 将段落转换为 HTML
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    try:
                        h_level = int(level)
                    except ValueError:
                        h_level = 2
                    content_parts.append(f"<h{h_level}>{text}</h{h_level}>")
                else:
                    content_parts.append(f"<p>{text}</p>")

        content = "\n".join(content_parts)

        # 检查是否已存在
        existing = db.query(Law).filter(Law.title == title).first()
        if existing:
            print(f"法规已存在: {title} (ID: {existing.id})")
            return

        # 创建法规记录
        law_data = {
            "title": title,
            "category": category,
            "publish_date": date.today(),
            "content": content,
            "source_url": "internal://local",  # 内部法规标记
            "is_internal": 1,
        }

        law = create_law(db, law_data)
        print(f"成功导入: {title} (ID: {law.id})")

    except Exception as e:
        print(f"导入失败: {e}")
        import traceback
        traceback.print_exc()
    finally:
        db.close()


def main():
    """主函数"""
    # 默认导入文件
    default_file = "/Users/ali/dev/usefulfiles/科研采购外协管理办法.docx"

    # 从命令行参数获取文件路径
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = default_file

    if not Path(file_path).exists():
        print(f"文件不存在: {file_path}")
        sys.exit(1)

    print(f"正在导入: {file_path}")
    import_docx(file_path)


if __name__ == "__main__":
    main()
